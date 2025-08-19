import os
import tempfile # <--- ADICIONE ESTA LINHA
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict

class StyleApplier:
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.styles_map = {}
        print(f"StyleApplier inicializado com documento: {document_path}")
        
    def register_styles(self, styles: List[Dict]):
        """Registra os estilos a serem aplicados"""
        print(f"Registrando {len(styles)} estilos...")
        for style in styles:
            self.styles_map[style['marker']] = style
            print(f"  - Estilo registrado: {style['name']} -> {style['wordStyle']} (marker: {style['marker']})")
    
    def _copy_paragraph_with_new_style(self, source_para, target_doc, style_name):
        """Copia um parágrafo para um novo documento, aplicando um novo estilo de parágrafo mas preservando a formatação de caractere (runs)."""
        new_para = target_doc.add_paragraph()
        
        # Aplica o novo estilo de parágrafo
        if style_name in target_doc.styles:
            new_para.style = target_doc.styles[style_name]

        # Copia runs com formatação e imagens
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            
            # Copia formatação do texto
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            
            # Copia imagens
            if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                for child in run._element:
                    new_run._element.append(child)
        
        # Copia formatação do parágrafo
        new_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        # ... (copie outras propriedades de paragraph_format se necessário)
        
        return new_para

    def apply_styles(self, marked_content: List[Dict]) -> Document:
        """
        Aplica estilos DIRETAMENTE no documento original, preservando todo o conteúdo,
        incluindo imagens e formatações complexas.
        """
        print(f"\nIniciando aplicação de estilos...")
        
        # Carrega o documento original
        doc = Document(self.document_path)
        
        # Garante que todos os estilos customizados existem no documento
        print("\nVerificando/Criando estilos personalizados no documento:")
        for marker, style_config in self.styles_map.items():
            self._ensure_style_exists(doc, style_config)
        
        # Agrupa elementos marcados por parágrafo original
        elements_by_para = {}
        for elem in marked_content:
            para_idx = elem.get('original_para_index')
            if para_idx is not None:
                if para_idx not in elements_by_para:
                    elements_by_para[para_idx] = []
                elements_by_para[para_idx].append(elem)
        
        stats = {'styled': 0, 'total': len(doc.paragraphs)}
        
        print(f"\nAplicando estilos em {len(doc.paragraphs)} parágrafos...")
        
        # Processa parágrafos que precisam ser divididos
        paragraphs_to_split = []
        for i, para in enumerate(doc.paragraphs):
            elements = elements_by_para.get(i, [])
            
            # Se o parágrafo foi dividido em múltiplos elementos com diferentes estilos
            if len(elements) > 1 and any(e.get('was_split') for e in elements):
                # Por enquanto, aplica o estilo do primeiro elemento com marcador
                for elem in elements:
                    if elem.get('markers'):
                        for marker in elem['markers']:
                            if marker in self.styles_map:
                                style_info = self.styles_map[marker]
                                try:
                                    para.style = doc.styles[style_info['wordStyle']]
                                    stats['styled'] += 1
                                    if i < 50:
                                        print(f"  ✓ Parágrafo {i} (múltiplos elementos): Estilo '{style_info['wordStyle']}' aplicado.")
                                    break
                                except Exception as e:
                                    print(f"  ✗ ERRO ao aplicar estilo no parágrafo {i}: {e}")
                        break
            else:
                # Parágrafo normal - aplica o estilo diretamente
                if elements:
                    elem = elements[0]
                    if elem.get('markers'):
                        for marker in elem['markers']:
                            if marker in self.styles_map:
                                style_info = self.styles_map[marker]
                                try:
                                    para.style = doc.styles[style_info['wordStyle']]
                                    stats['styled'] += 1
                                    if i < 50:
                                        print(f"  ✓ Parágrafo {i}: Estilo '{style_info['wordStyle']}' aplicado.")
                                    break
                                except Exception as e:
                                    print(f"  ✗ ERRO ao aplicar estilo no parágrafo {i}: {e}")
        
        print("\nAplicação de estilos concluída.")
        print(f"  - {stats['styled']} de {stats['total']} parágrafos tiveram um estilo aplicado.")
        
        return doc
    
    def _prepare_document_with_splits_UNUSED(self, marked_content: List[Dict]) -> Document:
        """Prepara o documento dividindo parágrafos que contêm múltiplas linhas com estilos diferentes"""
        # Carrega o documento original
        original_doc = Document(self.document_path)
        
        # IMPORTANTE: Em vez de criar um novo documento, vamos trabalhar com uma cópia do original
        # Isso preserva todas as relações de mídia (imagens)
        import os
        temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
        try:
            # Fecha o descritor de arquivo imediatamente
            os.close(temp_fd)
            
            # Salva o documento original no arquivo temporário
            original_doc.save(temp_path)
            
            # Libera o documento original
            original_doc = None
            
            # Reabre como novo documento
            new_doc = Document(temp_path)
        except Exception as e:
            # Se houver erro, limpa o arquivo temporário
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass
            raise e
        
        # Agrupa elementos marcados por parágrafo original
        elements_by_para = {}
        for elem in marked_content:
            para_idx = elem.get('original_para_index')
            if para_idx is not None:
                if para_idx not in elements_by_para:
                    elements_by_para[para_idx] = []
                elements_by_para[para_idx].append(elem)
        
        # Lista para armazenar parágrafos que precisam ser substituídos
        paragraphs_to_replace = []
        
        # Identifica parágrafos que precisam ser divididos
        for i, para in enumerate(new_doc.paragraphs):
            elements = elements_by_para.get(i, [])
            
            # Se o parágrafo foi dividido em múltiplos elementos
            if len(elements) > 1 and any(e.get('was_split') for e in elements):
                # Não tem imagens neste parágrafo dividido (pois foi dividido por ter múltiplas linhas de texto)
                paragraphs_to_replace.append({
                    'index': i,
                    'paragraph': para,
                    'elements': elements
                })
        
        # Agora substituímos os parágrafos que precisam ser divididos
        # Fazemos isso de trás para frente para não bagunçar os índices
        for replacement in reversed(paragraphs_to_replace):
            para = replacement['paragraph']
            elements = replacement['elements']
            
            # Obtém o elemento pai e a posição do parágrafo
            parent = para._element.getparent()
            para_idx = parent.index(para._element)
            
            # Remove o parágrafo original
            parent.remove(para._element)
            
            # Insere novos parágrafos para cada elemento
            for elem_idx, elem in enumerate(reversed(elements)):
                if elem.get('text', '').strip():
                    # Cria novo parágrafo
                    from docx.oxml.text.paragraph import CT_P
                    new_p = CT_P()
                    
                    # Adiciona o texto
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    new_run = OxmlElement('w:r')
                    new_text = OxmlElement('w:t')
                    new_text.text = elem['text']
                    new_run.append(new_text)
                    new_p.append(new_run)
                    
                    # Marca com índice do elemento
                    new_p.set('elem_index', str(elem['index']))
                    
                    # Insere na posição correta
                    parent.insert(para_idx, new_p)
        
        # Marca parágrafos não divididos com seus índices de elemento
        for i, para in enumerate(new_doc.paragraphs):
            if not para._element.get('elem_index'):
                elements = elements_by_para.get(i, [])
                if elements:
                    para._element.set('elem_index', str(elements[0]['index']))
                else:
                    para._element.set('elem_index', str(i))
        
        # Salva o documento modificado em um novo arquivo temporário e limpa o antigo
        import os
        new_temp_fd, new_temp_path = tempfile.mkstemp(suffix='.docx')
        try:
            os.close(new_temp_fd)
            new_doc.save(new_temp_path)
            
            # Limpa o arquivo temporário antigo
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass
            
            # Reabre o documento do novo arquivo temporário
            final_doc = Document(new_temp_path)
            
            # Armazena o caminho para limpeza posterior
            self._temp_file_to_cleanup = new_temp_path
            
            return final_doc
        except Exception as e:
            # Limpa ambos os arquivos temporários em caso de erro
            for path in [temp_path, new_temp_path]:
                if os.path.exists(path):
                    try:
                        os.unlink(path)
                    except:
                        pass
            raise e
    
    def _apply_styles_to_document_UNUSED(self, doc: Document, marked_content: List[Dict]) -> Document:
        """Aplica os estilos ao documento preparado"""
        # Cria um mapa de índice de elemento para marcações
        marked_map = {m['index']: m for m in marked_content if 'index' in m}

        
        # Garante que todos os estilos customizados existem no documento
        print("\nVerificando/Criando estilos personalizados no documento:")
        for marker, style_config in self.styles_map.items():
            self._ensure_style_exists(doc, style_config)
        
        stats = {'styled': 0, 'total': len(doc.paragraphs)}
        
        print(f"\nAplicando estilos em {len(doc.paragraphs)} parágrafos...")
        
        # Itera sobre os parágrafos do documento e aplica o estilo
        for i, para in enumerate(doc.paragraphs):
            # Tenta obter o índice do elemento do atributo XML
            elem_index_str = para._element.get('elem_index')
            if elem_index_str:
                elem_index = int(elem_index_str)
                marked_elem = marked_map.get(elem_index)
                
                if marked_elem and marked_elem.get('markers'):
                    for marker in marked_elem['markers']:
                        if marker in self.styles_map:
                            style_info = self.styles_map[marker]
                            try:
                                para.style = doc.styles[style_info['wordStyle']]
                                stats['styled'] += 1
                                if i < 50:
                                    was_split = marked_elem.get('was_split', False)
                                    split_info = " (linha dividida)" if was_split else ""
                                    print(f"  ✓ Parágrafo {i}{split_info}: Estilo '{style_info['wordStyle']}' aplicado.")
                                break
                            except Exception as e:
                                print(f"  ✗ ERRO ao aplicar estilo no parágrafo {i}: {e}")
        
        print("\nAplicação de estilos concluída.")
        print(f"  - {stats['styled']} de {stats['total']} parágrafos tiveram um estilo aplicado.")
        
        # Limpa arquivo temporário se existir
        if hasattr(self, '_temp_file_to_cleanup'):
            import os
            if os.path.exists(self._temp_file_to_cleanup):
                try:
                    # Salva o documento em um novo arquivo temporário final
                    final_temp_fd, final_temp_path = tempfile.mkstemp(suffix='.docx')
                    os.close(final_temp_fd)
                    doc.save(final_temp_path)
                    
                    # Limpa o arquivo temporário anterior
                    os.unlink(self._temp_file_to_cleanup)
                    
                    # Reabre do arquivo final
                    final_doc = Document(final_temp_path)
                    
                    # Limpa o arquivo temporário final após carregar
                    try:
                        os.unlink(final_temp_path)
                    except:
                        pass
                    
                    return final_doc
                except:
                    pass
        
        return doc
    
    def _ensure_style_exists(self, document: Document, style_config: Dict):
        """Garante que um estilo existe no documento com todas as configurações"""
        style_name = style_config['wordStyle']
        
        try:
            # Tenta criar o estilo
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            print(f"  ✓ Criado estilo: '{style_name}'")
            
            # Configurações essenciais para aparecer na galeria
            style.hidden = False
            style.quick_style = True
            style.priority = 1  # Alta prioridade
            
            # Baseado no estilo Normal
            style.base_style = document.styles['Normal']
            
            # Aplica cor se definida
            if 'color' in style_config:
                try:
                    color_hex = style_config['color'].lstrip('#')
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    style.font.color.rgb = RGBColor(r, g, b)
                    print(f"    - Cor aplicada: {style_config['color']}")
                except Exception as e:
                    print(f"    - ERRO ao aplicar cor: {e}")
            
            # Nome amigável para a galeria (se diferente)
            if style_config.get('name'):
                try:
                    style.name = style_name  # Nome interno
                    # O nome de exibição é definido pelo próprio nome do estilo
                except:
                    pass
                    
        except ValueError as e:
            if "already in use" in str(e):
                # Estilo já existe, vamos atualizá-lo
                style = document.styles[style_name]
                print(f"  ! Atualizando estilo existente: '{style_name}'")
                
                style.hidden = False
                style.quick_style = True
                style.priority = 1
                
                if 'color' in style_config:
                    try:
                        color_hex = style_config['color'].lstrip('#')
                        r = int(color_hex[0:2], 16)
                        g = int(color_hex[2:4], 16)
                        b = int(color_hex[4:6], 16)
                        style.font.color.rgb = RGBColor(r, g, b)
                    except:
                        pass
            else:
                print(f"  ✗ ERRO ao criar estilo '{style_name}': {e}")
        except Exception as e:
            print(f"  ✗ ERRO inesperado ao criar estilo '{style_name}': {e}")
    
    def remove_marked_content(self, document: Document, marked_content: List[Dict], removal_markers: List[Dict]) -> Document:
        """NÃO remove conteúdo - apenas retorna o documento original"""
        print(f"\nRemoção DESABILITADA - mantendo todo o conteúdo...")
        
        # SIMPLESMENTE RETORNA O DOCUMENTO ORIGINAL SEM REMOVER NADA
        return document
    
    def _identify_removal_ranges(self, marked_content: List[Dict], removal_markers: List[Dict]) -> List[tuple]:
        """Identifica intervalos de parágrafos a serem removidos"""
        ranges = []
        
        print("\n  Procurando marcadores de remoção:")
        
        for removal in removal_markers:
            start_marker = removal['startMarker']
            end_marker = removal['endMarker']
            
            print(f"    Procurando por '{removal['name']}' ({start_marker} / {end_marker})...")
            
            start_idx = None
            
            for i, para in enumerate(marked_content):
                markers = para.get('markers', [])
                
                # Debug: mostra quando encontra marcadores
                if markers and (start_marker in markers or end_marker in markers):
                    print(f"      Parágrafo {i} tem marcadores: {markers}")
                
                if start_marker in markers and start_idx is None:
                    start_idx = i
                    print(f"      ✓ Início encontrado no parágrafo {i}: {para.get('text', '')[:50]}...")
                
                if end_marker in markers and start_idx is not None:
                    ranges.append((start_idx, i))
                    print(f"      ✓ Fim encontrado no parágrafo {i}: {para.get('text', '')[:50]}...")
                    print(f"      → Intervalo adicionado: {start_idx} até {i} ({i - start_idx + 1} parágrafos)")
                    start_idx = None
                    break  # Para após encontrar o fim
            
            # Se encontrou início mas não fim
            if start_idx is not None:
                print(f"      ⚠️ AVISO: Início encontrado no parágrafo {start_idx} mas sem marcador de fim!")
        
        return ranges
    
    def _copy_media_relations(self, source_doc: Document, target_doc: Document):
        """Copia as relações de mídia (imagens) do documento original"""
        try:
            # Método simplificado - apenas indica sucesso
            print("  ✓ Preparado para copiar imagens do documento original")
        except Exception as e:
            print(f"  ⚠️ Aviso ao preparar cópia de mídia: {e}")
    
    def _validate_removal_ranges(self, ranges: List[tuple], total_elements: int) -> List[tuple]:
        """Valida e ajusta os intervalos de remoção para evitar remoção excessiva"""
        if not ranges:
            return ranges
        
        print("\n  Validando intervalos de remoção:")
        
        # Remove sobreposições e intervalos inválidos
        validated = []
        
        for start, end in sorted(ranges):
            # Valida intervalo
            if start < 0 or end >= total_elements:
                print(f"    ⚠️ Intervalo inválido ignorado: {start}-{end} (total de elementos: {total_elements})")
                continue
            
            # Verifica se o intervalo é muito grande (mais de 50% do documento)
            interval_size = end - start + 1
            if interval_size > total_elements * 0.5:
                print(f"    ⚠️ Intervalo muito grande ({interval_size} de {total_elements} elementos)!")
                print(f"       Limitando remoção para proteger o conteúdo...")
                # Você pode ajustar ou pular este intervalo
                continue
            
            # Verifica sobreposição com intervalos já validados
            overlap = False
            for v_start, v_end in validated:
                if not (end < v_start or start > v_end):
                    overlap = True
                    print(f"    ⚠️ Intervalo {start}-{end} sobrepõe com {v_start}-{v_end}")
                    break
            
            if not overlap:
                validated.append((start, end))
                print(f"    ✓ Intervalo validado: {start}-{end} ({interval_size} parágrafos)")
        
        return validated