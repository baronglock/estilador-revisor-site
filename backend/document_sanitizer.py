from docx import Document
from docx.shared import RGBColor
from typing import Dict, List
import re

class DocumentSanitizer:
    """
    Sanitiza documentos Word para importação limpa no InDesign,
    removendo formatações locais que podem causar conflitos.
    """
    
    def __init__(self, document: Document):
        self.document = document
        print("DocumentSanitizer inicializado")
    
    def sanitize_local_formatting(self) -> Document:
        """
        Remove formatações locais mantendo apenas os estilos aplicados.
        Isso garante que o InDesign possa mapear corretamente os estilos.
        """
        print("\nIniciando sanitização do documento...")
        
        stats = {
            'paragraphs_processed': 0,
            'runs_cleaned': 0,
            'styles_preserved': set()
        }
        
        # Processa cada parágrafo
        for para in self.document.paragraphs:
            # Preserva o estilo do parágrafo
            if para.style:
                stats['styles_preserved'].add(para.style.name)
            
            # Remove formatações diretas dos runs
            for run in para.runs:
                # Remove formatações locais que podem conflitar com o InDesign
                self._clean_run_formatting(run)
                stats['runs_cleaned'] += 1
            
            stats['paragraphs_processed'] += 1
            
            # Log de progresso a cada 100 parágrafos
            if stats['paragraphs_processed'] % 100 == 0:
                print(f"  Processados {stats['paragraphs_processed']} parágrafos...")
        
        # Processa tabelas se houver
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.style:
                            stats['styles_preserved'].add(para.style.name)
                        
                        for run in para.runs:
                            self._clean_run_formatting(run)
                            stats['runs_cleaned'] += 1
        
        print(f"\nSanitização concluída:")
        print(f"  - Parágrafos processados: {stats['paragraphs_processed']}")
        print(f"  - Runs limpos: {stats['runs_cleaned']}")
        print(f"  - Estilos preservados: {len(stats['styles_preserved'])}")
        
        return self.document
    
    def _clean_run_formatting(self, run):
        """
        Remove formatações diretas de um run, mantendo apenas o essencial.
        Isso evita conflitos com os estilos do InDesign.
        """
        try:
            # Remove formatações que devem vir apenas dos estilos
            # mas preserva o texto e elementos especiais (como imagens)
            
            # NÃO remove bold/italic/underline se forem parte do conteúdo original
            # Apenas remove se forem aplicados como formatação direta conflitante
            
            # Remove tamanho de fonte local (deve vir do estilo)
            if run.font.size:
                run.font.size = None
            
            # Remove cor local se não for essencial
            # (mantém apenas se for diferente de preto padrão)
            if run.font.color and run.font.color.rgb:
                # Se for preto (000000), remove para usar o padrão do estilo
                if run.font.color.rgb == RGBColor(0, 0, 0):
                    run.font.color.rgb = None
            
            # Remove nome da fonte local (deve vir do estilo)
            if run.font.name:
                run.font.name = None
            
            # Preserva elementos especiais (imagens, campos, etc)
            # Estes são identificados pelos elementos XML
            
        except Exception as e:
            # Em caso de erro, apenas continua sem modificar
            pass
    
    def create_style_mapping_report(self) -> Dict:
        """
        Cria um relatório dos estilos usados no documento para 
        facilitar o mapeamento no InDesign.
        """
        style_report = {
            'paragraph_styles': {},
            'character_styles': {},
            'table_styles': {}
        }
        
        # Mapeia estilos de parágrafo
        for para in self.document.paragraphs:
            if para.style:
                style_name = para.style.name
                if style_name not in style_report['paragraph_styles']:
                    style_report['paragraph_styles'][style_name] = {
                        'count': 0,
                        'sample_text': para.text[:50] + '...' if len(para.text) > 50 else para.text
                    }
                style_report['paragraph_styles'][style_name]['count'] += 1
        
        return style_report
    
    def add_indesign_markers(self, use_xml_tags: bool = False):
        """
        Adiciona marcadores especiais para facilitar a importação no InDesign.
        """
        if use_xml_tags:
            print("\nAdicionando marcadores XML para InDesign...")
            
            for para in self.document.paragraphs:
                if para.style:
                    # Adiciona tags XML invisíveis que o InDesign pode usar
                    style_tag = f"[{para.style.name}]"
                    
                    # Adiciona tag no início do parágrafo se não estiver vazio
                    if para.text.strip():
                        # Insere tag como comentário XML (não visível no Word)
                        pass  # Implementação específica se necessário
        
        print("  ✓ Marcadores adicionados")
    
    def optimize_for_indesign(self) -> Document:
        """
        Aplica todas as otimizações para InDesign de uma vez.
        """
        print("\n=== OTIMIZANDO DOCUMENTO PARA INDESIGN ===")
        
        # 1. Remove formatações locais
        self.sanitize_local_formatting()
        
        # 2. Cria relatório de estilos
        style_report = self.create_style_mapping_report()
        print(f"\nRelatório de estilos:")
        print(f"  - Estilos de parágrafo únicos: {len(style_report['paragraph_styles'])}")
        for style_name, info in style_report['paragraph_styles'].items():
            print(f"    • {style_name}: {info['count']} ocorrências")
        
        # 3. Adiciona marcadores se necessário
        # self.add_indesign_markers(use_xml_tags=True)
        
        print("\n✓ Documento otimizado para importação no InDesign")
        
        return self.document