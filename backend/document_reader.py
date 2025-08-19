from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re

class DocumentReader:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        
    def read_paragraphs(self):
        """Lê todos os parágrafos e elementos do documento incluindo imagens"""
        elements = []
        element_index = 0
        
        # Primeiro, processa parágrafos normais
        for i, para in enumerate(self.document.paragraphs):
            # Verifica se o parágrafo contém imagem inline
            has_inline_image = False
            for run in para.runs:
                if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                    has_inline_image = True
                    print(f"  Imagem inline detectada no parágrafo {i}")
                    break
            is_image_paragraph = has_inline_image and not para.text.strip()
            
            # Verifica se o parágrafo tem múltiplas linhas que deveriam ser elementos separados
            para_text = para.text
            if para_text and '\n' in para_text:
                lines = para_text.split('\n')
                
                # Detecta se as linhas são de tipos diferentes (questão vs alternativas)
                should_split = self._should_split_paragraph_lines(lines)
                
                if should_split:
                    print(f"  Parágrafo {i} será dividido em {len(lines)} elementos separados")
                    
                    # Cria um elemento para cada linha significativa
                    for line_idx, line in enumerate(lines):
                        if not line.strip():
                            continue
                        
                        # Detecção de listas para cada linha
                        is_list_item, list_type, list_char = self._detect_list_item(line.strip())
                        
                        elements.append({
                            'index': element_index,
                            'type': 'paragraph',
                            'text': line,
                            'original_para_index': i,
                            'line_in_paragraph': line_idx,
                            'was_split': True,
                            'style': para.style.name if para.style else 'Normal',
                            'runs': self._extract_runs_for_line(para, line),
                            'has_image': False,
                            'is_image_paragraph': False,
                            'is_list_item': is_list_item,
                            'list_type': list_type,
                            'list_char': list_char,
                            'markers': []
                        })
                        element_index += 1
                    continue
            
            # Detecção detalhada de listas
            is_list_item = False
            list_type = None
            list_char = None
            
            # Verifica se é um item de lista formatado pelo Word
            if para._element.xpath('.//w:numPr'):
                is_list_item = True
                
                # Tenta identificar o tipo baseado no texto
                text_start = para.text.strip()[:10] if para.text else ""
                
                # Detecta tipo de marcador
                if text_start:
                    # Lista com letras (a), b), A), B)
                    if re.match(r'^[a-zA-Z][\)\.]\s', text_start):
                        list_type = 'letter'
                        list_char = text_start[0]
                    # Lista numerada 1. 2. 1) 2)
                    elif re.match(r'^\d+[\)\.]\s', text_start):
                        list_type = 'number'
                    # Bullets (•, -, *, etc)
                    else:
                        list_type = 'bullet'
                        # Tenta identificar o caractere do bullet
                        if para._element.xpath('.//w:lvlText'):
                            list_char = 'bullet'
            
            # Verifica também manualmente se parece uma lista (caso não esteja formatada)
            elif para.text:
                text_start = para.text.strip()
                # Padrões manuais de lista
                if re.match(r'^[a-eA-E][\)\.]\s', text_start):
                    is_list_item = True
                    list_type = 'letter'
                    list_char = text_start[0].upper()
                elif re.match(r'^\d+[\)\.]\s', text_start):
                    is_list_item = True
                    list_type = 'number'
                elif text_start.startswith(('• ', '- ', '* ', '→ ', '▪ ')):
                    is_list_item = True
                    list_type = 'bullet'
                    list_char = text_start[0]
            
            # SEMPRE adiciona o parágrafo, mesmo se vazio
            elements.append({
                'index': element_index,
                'type': 'paragraph',
                'text': para.text,  # Pode ser vazio
                'original_para_index': i,
                'style': para.style.name if para.style else 'Normal',
                'runs': self._extract_runs(para),
                'has_image': has_inline_image,
                'is_image_paragraph': is_image_paragraph, # <--- LINHA ADICIONADA/MODIFICADA
                'is_list_item': is_list_item,
                'list_type': list_type,
                'list_char': list_char,
                'markers': []
            })
            element_index += 1
        
        # Processa tabelas
        for i, table in enumerate(self.document.tables):
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                elements.append({
                    'index': element_index,
                    'type': 'table',
                    'text': '\n'.join(table_text),
                    'original_element': table._element,
                    'style': 'Table',
                    'markers': []
                })
                element_index += 1
        
        print(f"\nTotal de elementos lidos: {len(elements)}")
        
        # Conta tipos de elementos
        types_count = {}
        for elem in elements:
            elem_type = elem['type']
            types_count[elem_type] = types_count.get(elem_type, 0) + 1
        
        for elem_type, count in types_count.items():
            print(f"  - {elem_type}: {count}")
        
        # Debug: mostra onde estão as imagens
        image_positions = [elem['index'] for elem in elements if elem['type'] == 'image']
        if image_positions:
            print(f"  Posições das imagens: {image_positions}")
        
        return elements
    
    def _extract_runs(self, paragraph):
        """Extrai informações de formatação dos runs"""
        runs = []
        for run in paragraph.runs:
            runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            })
        return runs
    
    def _should_split_paragraph_lines(self, lines):
        """Determina se as linhas de um parágrafo devem ser separadas em elementos distintos"""
        if len(lines) <= 1:
            return False
        
        # IMPORTANTE: Sempre separar se há múltiplas linhas com conteúdo significativo
        # Isso garante que cada linha possa receber seu próprio estilo
        non_empty_lines = [line for line in lines if line.strip()]
        
        # Se tem 2 ou mais linhas não vazias, deve verificar se precisa separar
        if len(non_empty_lines) < 2:
            return False
        
        # Padrões que indicam diferentes tipos de conteúdo
        patterns = {
            'question': [
                r'^\d+[\.\)]\s',  # Numeração de questão (1. ou 1))
                r'^\d+\s*[-–]\s',  # 1 - 
                r'^Questão\s+\d+',  # "Questão 1"
                r'^Q\d+[\.\)]\s',  # Q1. ou Q1)
                r'^QUESTÃO\s+\d+',  # QUESTÃO 1
            ],
            'alternative': [
                r'^[a-eA-E][\.\)]\s',  # a) b) c) d) e) ou a. b. c. etc
                r'^\([a-eA-E]\)',  # (a) (b) (c) etc
                r'^[A-E]\s*[-–]\s',  # A - B - C - etc
                r'^[a-eA-E]\s',  # Apenas letra seguida de espaço
            ],
            'answer': [
                r'^Resposta:',
                r'^Gabarito:',
                r'^Alternativa correta:',
                r'^[a-eA-E]\d+\s*[-–]',  # a1- b2- etc (padrão de gabarito)
                r'^GABARITO',
            ],
            'title': [
                r'^Simulado\s+\d+',
                r'^SIMULADO\s+\d+',
                r'^Prova\s+\d+',
                r'^Teste\s+\d+',
            ],
            'subtitle': [
                r'^Estudos\s+\d+',
                r'^Parte\s+[IVX]+',
                r'^Seção\s+\d+',
            ]
        }
        
        # Analisa cada linha para determinar seu tipo
        line_types = []
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                line_types.append('empty')
                continue
            
            found_type = 'text'
            for pattern_type, pattern_list in patterns.items():
                for pattern in pattern_list:
                    if re.match(pattern, line_stripped, re.IGNORECASE):
                        found_type = pattern_type
                        break
                if found_type != 'text':
                    break
            
            line_types.append(found_type)
        
        # Remove linhas vazias para análise
        non_empty_types = [t for t in line_types if t != 'empty']
        
        # Se todas as linhas são do mesmo tipo, não precisa separar
        # EXCETO se são todas alternativas (cada alternativa deve ser separada)
        if len(set(non_empty_types)) == 1 and non_empty_types[0] != 'alternative':
            return False
        
        # Se há diferentes tipos, SEMPRE separa
        if len(set(non_empty_types)) > 1:
            return True
        
        # Se são todas alternativas mas há mais de uma, separa
        if non_empty_types and non_empty_types[0] == 'alternative' and len(non_empty_types) > 1:
            return True
        
        # Análise adicional: se as linhas começam de forma muito diferente, separa
        # Isso captura casos não cobertos pelos padrões acima
        if len(non_empty_lines) >= 2:
            first_words = []
            for line in non_empty_lines[:3]:  # Analisa até 3 primeiras linhas
                words = line.strip().split()
                if words:
                    first_words.append(words[0].lower())
            
            # Se os primeiros words são muito diferentes, provavelmente são conteúdos diferentes
            if len(set(first_words)) == len(first_words):
                # Todas diferentes - pode indicar necessidade de separação
                return True
        
        return False
    
    def _detect_list_item(self, text):
        """Detecta se um texto é um item de lista e retorna o tipo"""
        is_list_item = False
        list_type = None
        list_char = None
        
        # Padrões de lista
        if re.match(r'^[a-eA-E][\)\.]\s', text):
            is_list_item = True
            list_type = 'letter'
            list_char = text[0].upper()
        elif re.match(r'^\([a-eA-E]\)', text):
            is_list_item = True
            list_type = 'letter'
            list_char = text[1].upper()
        elif re.match(r'^\d+[\)\.]\s', text):
            is_list_item = True
            list_type = 'number'
        elif text.startswith(('• ', '- ', '* ', '→ ', '▪ ')):
            is_list_item = True
            list_type = 'bullet'
            list_char = text[0]
        
        return is_list_item, list_type, list_char
    
    def _extract_runs_for_line(self, paragraph, line_text):
        """Extrai runs específicos para uma linha de texto dentro de um parágrafo"""
        # Por enquanto, retorna os runs do parágrafo inteiro
        # Em uma implementação mais sofisticada, poderíamos mapear os runs para cada linha
        runs = []
        
        # Simplificação: usa a formatação do primeiro run não vazio
        for run in paragraph.runs:
            if run.text.strip():
                runs.append({
                    'text': line_text,  # Usa o texto da linha
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                })
                break
        
        if not runs:
            # Se não encontrou runs com texto, cria um padrão
            runs.append({
                'text': line_text,
                'bold': None,
                'italic': None,
                'underline': None,
                'font_size': None,
                'font_color': None
            })
        
        return runs
    
    def read_tables(self):
        """Lê todas as tabelas do documento"""
        tables = []
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append({
                'index': i,
                'data': table_data
            })
        return tables
    
    def get_document_info(self):
        """Retorna informações gerais do documento"""
        # Conta elementos
        paragraph_count = 0
        image_count = 0
        table_count = 0
        
        for element in self.document.element.body:
            if element.tag.endswith('p'):
                paragraph_count += 1
                # Verifica se tem imagem
                para = None
                for p in self.document.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    for run in para.runs:
                        if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                            image_count += 1
                            break
            elif element.tag.endswith('tbl'):
                table_count += 1
        
        return {
            'total_paragraphs': paragraph_count,
            'total_images': image_count,
            'total_tables': table_count,
            'total_sections': len(self.document.sections),
            'core_properties': {
                'author': self.document.core_properties.author,
                'created': str(self.document.core_properties.created) if self.document.core_properties.created else None,
                'modified': str(self.document.core_properties.modified) if self.document.core_properties.modified else None,
                'title': self.document.core_properties.title
            }
        }