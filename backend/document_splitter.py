from docx import Document
from typing import List, Dict, Tuple
import re

class DocumentSplitter:
    def __init__(self):
        self.simulado_pattern = re.compile(r'Simulado\s+(\d+)', re.IGNORECASE)
        
    def split_simulados(self, document: Document) -> List[Dict]:
        """Divide o documento em simulados individuais com precisão melhorada"""
        print("\n=== DIVIDINDO DOCUMENTO EM SIMULADOS ===")
        simulados = []
        current_simulado = None
        current_content = []
        current_start = 0
        
        # Primeira passada: identifica todos os títulos de simulado
        simulado_positions = []
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip()
            
            # Verifica se é título de simulado (mais preciso)
            if self._is_simulado_title(text, para):
                match = self.simulado_pattern.search(text)
                if match:
                    simulado_num = int(match.group(1))
                    simulado_positions.append({
                        'number': simulado_num,
                        'index': i,
                        'text': text
                    })
                    print(f"  Encontrado Simulado {simulado_num} na posição {i}: {text[:50]}...")
        
        # Segunda passada: coleta o conteúdo de cada simulado
        for idx, sim_pos in enumerate(simulado_positions):
            start_idx = sim_pos['index']
            end_idx = simulado_positions[idx + 1]['index'] if idx + 1 < len(simulado_positions) else len(document.paragraphs)
            
            # Coleta parágrafos do simulado
            content = []
            for i in range(start_idx, end_idx):
                content.append(document.paragraphs[i])
            
            simulados.append({
                'number': sim_pos['number'],
                'content': content,
                'start_index': start_idx,
                'end_index': end_idx - 1,
                'paragraph_count': len(content)
            })
            
            print(f"  Simulado {sim_pos['number']}: {len(content)} parágrafos (índices {start_idx}-{end_idx-1})")
        
        # Validação
        print(f"\nTotal de simulados encontrados: {len(simulados)}")
        total_paragraphs = sum(s['paragraph_count'] for s in simulados)
        print(f"Total de parágrafos em simulados: {total_paragraphs}")
        print(f"Total de parágrafos no documento: {len(document.paragraphs)}")
        
        if total_paragraphs < len(document.paragraphs):
            print(f"AVISO: {len(document.paragraphs) - total_paragraphs} parágrafos não atribuídos a nenhum simulado")
        
        return simulados
    
    def _is_simulado_title(self, text: str, paragraph) -> bool:
        """Verifica se o parágrafo é realmente um título de simulado"""
        # Verifica o padrão de texto
        if not self.simulado_pattern.search(text):
            return False
        
        # Verifica se é um título curto (evita falsos positivos em texto corrido)
        if len(text) > 100:
            return False
        
        # Verifica o estilo (títulos geralmente têm estilos específicos)
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if any(h in style_name for h in ['heading', 'título', 'title']):
                return True
        
        # Verifica formatação (negrito, tamanho maior)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                return True
        
        # Se tem apenas o padrão "Simulado X" ou similar, provavelmente é título
        if len(text.split()) <= 10:
            return True
        
        return False
    
    def separate_questions_answers_enhanced(self, simulado_content: List, 
                                          marked_content: List[Dict], 
                                          simulado_info: Dict,
                                          styles: List[Dict]) -> Tuple[List, List]:
        """Separa questões de gabaritos usando as marcações da IA e estilos definidos"""
        print(f"\n  Separando Simulado {simulado_info['number']} usando marcações...")
        
        questions_content = []
        answers_content = []
        
        # Mapeia os parágrafos originais para as marcações
        start_idx = simulado_info['start_index']
        end_idx = simulado_info['end_index']
        
        gabarito_count = 0
        question_count = 0
        
        for i, para in enumerate(simulado_content):
            # Encontra a marcação correspondente
            global_idx = start_idx + i
            
            # Busca nas marcações
            marked_para = None
            if global_idx < len(marked_content):
                marked_para = marked_content[global_idx]
            
            is_gabarito = False
            
            # Verifica marcações da IA primeiro
            if marked_para and marked_para.get('markers'):
                markers = marked_para['markers']
                # Verifica se tem marcador de gabarito baseado nos estilos definidos
                for marker in markers:
                    style_info = next((s for s in styles if s.get('marker') == marker), None)
                    if style_info:
                        style_name = style_info.get('name', '').lower()
                        if any(word in style_name for word in ['gabarito', 'resposta', 'answer']):
                            is_gabarito = True
                            break
            
            # Fallback: verifica pelo estilo original se não tem marcação
            if not is_gabarito and para.style and para.style.name:
                style_name_lower = para.style.name.lower()
                if 'gabarito' in style_name_lower or 'answer' in style_name_lower or 'resposta' in style_name_lower:
                    is_gabarito = True
            
            # Fallback final: análise de conteúdo baseada nos estilos definidos
            if not is_gabarito:
                text_lower = para.text.lower().strip()
                
                # Busca nos prompts dos estilos definidos se algum indica gabarito
                for style in styles:
                    style_name = style.get('name', '').lower()
                    style_prompt = style.get('prompt', '').lower()
                    
                    if any(word in style_name for word in ['gabarito', 'resposta', 'answer']):
                        # Verifica se o texto atual corresponde ao prompt deste estilo
                        if self._text_matches_prompt(text_lower, style_prompt):
                            is_gabarito = True
                            break
            
            # Adiciona ao grupo apropriado
            if is_gabarito:
                answers_content.append(para)
                gabarito_count += 1
            else:
                questions_content.append(para)
                question_count += 1
        
        print(f"    - Questões: {question_count} parágrafos")
        print(f"    - Gabaritos: {gabarito_count} parágrafos")
        print(f"    - Total: {question_count + gabarito_count} parágrafos")
        
        return questions_content, answers_content
    
    def create_split_documents(self, simulados: List[Dict], document: Document, 
                              marked_content: List[Dict], styles: List[Dict]) -> Dict[str, Document]:
        """Cria documentos separados usando as marcações da IA e estilos definidos"""
        print("\n=== CRIANDO DOCUMENTOS SEPARADOS ===")
        documents = {}
        
        for simulado in simulados:
            sim_num = simulado['number']
            print(f"\nProcessando Simulado {sim_num}...")
            
            # Separa usando marcações da IA e estilos
            questions_content, answers_content = self.separate_questions_answers_enhanced(
                simulado['content'], 
                marked_content,
                simulado,
                styles
            )
            
            # Cria documento de questões
            if questions_content:
                questions_doc = self._create_document_from_paragraphs(
                    questions_content, 
                    document,
                    f"Simulado {sim_num} - Questões"
                )
                documents[f'simulado_{sim_num}_questoes'] = questions_doc
                print(f"  ✓ Documento de questões criado: {len(questions_content)} parágrafos")
            
            # Cria documento de gabaritos
            if answers_content:
                answers_doc = self._create_document_from_paragraphs(
                    answers_content, 
                    document,
                    f"Simulado {sim_num} - Gabarito"
                )
                documents[f'simulado_{sim_num}_gabarito'] = answers_doc
                print(f"  ✓ Documento de gabarito criado: {len(answers_content)} parágrafos")
        
        return documents
    
    def create_complete_documents(self, document: Document, marked_content: List[Dict]) -> Dict[str, Document]:
        """Cria documento completo e separados (questões/gabaritos) do documento inteiro"""
        print("\n=== CRIANDO DOCUMENTOS COMPLETOS ===")
        documents = {}
        
        # 1. Documento completo estilizado
        print("\nCriando documento completo...")
        complete_doc = Document()
        self._copy_styles(document, complete_doc)
        
        para_count = 0
        for para in document.paragraphs:
            if para.text.strip() or self._has_image(para):
                self._copy_paragraph_to_document(complete_doc, para)
                para_count += 1
        
        documents['completo'] = complete_doc
        print(f"  ✓ Documento completo: {para_count} parágrafos")
        
        # 2. Documento só com questões (todo o documento)
        print("\nCriando documento de todas as questões...")
        all_questions_doc = Document()
        all_answers_doc = Document()
        self._copy_styles(document, all_questions_doc)
        self._copy_styles(document, all_answers_doc)
        
        questions_count = 0
        answers_count = 0
        
        for i, para in enumerate(document.paragraphs):
            if i < len(marked_content):
                marked = marked_content[i]
                markers = marked.get('markers', [])
                
                # Verifica se é gabarito
                is_gabarito = any('GABARITO' in m for m in markers)
                
                if not is_gabarito:
                    # Também verifica padrões de gabarito no texto
                    text_lower = para.text.lower().strip()
                    gabarito_patterns = [
                        r'^[a-h]\d+\s*[–\-]',
                        r'^resposta:',
                        r'^gabarito:',
                        r'^alternativa correta:'
                    ]
                    for pattern in gabarito_patterns:
                        if re.search(pattern, text_lower):
                            is_gabarito = True
                            break
                
                if is_gabarito:
                    self._copy_paragraph_to_document(all_answers_doc, para)
                    answers_count += 1
                else:
                    self._copy_paragraph_to_document(all_questions_doc, para)
                    questions_count += 1
            else:
                # Se não tem marcação, adiciona às questões por padrão
                self._copy_paragraph_to_document(all_questions_doc, para)
                questions_count += 1
        
        documents['todas_questoes'] = all_questions_doc
        documents['todos_gabaritos'] = all_answers_doc
        
        print(f"  ✓ Documento de questões: {questions_count} parágrafos")
        print(f"  ✓ Documento de gabaritos: {answers_count} parágrafos")
        
        return documents
    
    def _create_document_from_paragraphs(self, paragraphs: List, source_doc: Document, 
                                       title: str = None) -> Document:
        """Cria um novo documento a partir de uma lista de parágrafos"""
        new_doc = Document()
        
        # Copia estilos
        self._copy_styles(source_doc, new_doc)
        
        # Adiciona título se fornecido
        if title:
            title_para = new_doc.add_paragraph(title)
            if 'Heading 1' in [s.name for s in new_doc.styles]:
                title_para.style = 'Heading 1'
            new_doc.add_paragraph()  # Linha em branco
        
        # Copia parágrafos
        for para in paragraphs:
            if para.text.strip() or self._has_image(para):
                self._copy_paragraph_to_document(new_doc, para)
        
        return new_doc
    
    def _has_image(self, paragraph) -> bool:
        """Verifica se o parágrafo contém imagem"""
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                return True
        return False
    
    def _text_matches_prompt(self, text: str, prompt: str) -> bool:
        """Verifica se o texto corresponde ao prompt definido pelo usuário"""
        # Extrai palavras-chave do prompt
        keywords = re.findall(r'\b\w+\b', prompt.lower())
        
        # Conta quantas palavras-chave estão presentes no texto
        matches = sum(1 for keyword in keywords if keyword in text)
        
        # Retorna True se tem correspondência significativa
        return matches >= min(2, len(keywords) // 2)
    
    def _copy_styles(self, source_doc: Document, target_doc: Document):
        """Copia estilos de um documento para outro com tratamento melhorado"""
        copied_styles = set()
        
        for style in source_doc.styles:
            try:
                if hasattr(style, 'name') and style.name not in copied_styles:
                    # Pula estilos built-in que já existem
                    if style.builtin:
                        continue
                    
                    # Tenta adicionar o estilo
                    try:
                        target_style = target_doc.styles.add_style(
                            style.name, 
                            style.type
                        )
                        
                        # Copia propriedades
                        target_style.hidden = style.hidden
                        target_style.quick_style = style.quick_style
                        
                        if hasattr(style, 'priority'):
                            target_style.priority = style.priority
                        
                        # Copia formatação
                        if hasattr(style, 'font'):
                            if style.font.color and style.font.color.rgb:
                                target_style.font.color.rgb = style.font.color.rgb
                            if style.font.size:
                                target_style.font.size = style.font.size
                            target_style.font.bold = style.font.bold
                            target_style.font.italic = style.font.italic
                        
                        copied_styles.add(style.name)
                        
                    except ValueError:
                        # Estilo já existe
                        pass
                        
            except Exception as e:
                # Ignora erros individuais
                pass
    
    def _copy_paragraph_to_document(self, target_doc: Document, source_para):
        """Copia um parágrafo mantendo toda formatação e conteúdo"""
        # Cria novo parágrafo
        new_para = target_doc.add_paragraph()
        
        # Copia o estilo
        try:
            if source_para.style:
                new_para.style = source_para.style.name
        except:
            pass
        
        # Copia runs com formatação e imagens
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            
            # Copia formatação do texto
            if run.bold is not None:
                new_run.bold = run.bold
            if run.italic is not None:
                new_run.italic = run.italic
            if run.underline is not None:
                new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            
            # Copia imagens se existirem
            if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                # Copia o elemento XML da imagem diretamente
                for child in run._element:
                    new_run._element.append(child)
        
        # Copia formatação do parágrafo
        if source_para.paragraph_format.alignment:
            new_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        if source_para.paragraph_format.left_indent:
            new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        if source_para.paragraph_format.right_indent:
            new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        if source_para.paragraph_format.first_line_indent:
            new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        if source_para.paragraph_format.space_before:
            new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        if source_para.paragraph_format.space_after:
            new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        if source_para.paragraph_format.line_spacing:
            new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing